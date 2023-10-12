import textwrap
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FileWriter:

    def write_text_to_pdf(filename, text):
        # save FPDF() class into
        # a variable pdf
        pdf = FPDF()

        # Add a page
        pdf.add_page()

        # set style and size of font
        # that you want in the pdf
        pdf.set_font("Arial", size=14)

        # Add the entire text as a multi-line block
        pdf.multi_cell(200, 10, txt=text, align='W')

        # save the pdf with name .pdf
        pdf.output(filename)

    def write_text_to_docx(company_name, overview_text, business_model_text, risk_factors_text, page_start, page_end):
        document = Document()

        for section in document.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        document.add_heading(company_name, 0)

        document.add_heading('Overview', level=1)

        p = document.add_paragraph(overview_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        document.add_heading('Our Business model', level=1)

        p = document.add_paragraph(business_model_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        document.add_heading('Financial Information', level=1)
        p = document.add_paragraph(
            'Here are screenshots which represent financial information.')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        new_section = document.add_section()
        new_section.top_margin = Cm(0)
        new_section.left_margin = Cm(0)
        new_section.right_margin = Cm(0)
        new_section.bottom_margin = Cm(0)

        for page in range(page_start, page_end):
            document.add_picture(f'{company_name} {page}.png', width=Cm(22))

        new_section = document.add_section()
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        document.add_heading('Risk Factors', level=1)
        p = document.add_paragraph(risk_factors_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        document.save(f'{company_name}_summary.docx')
